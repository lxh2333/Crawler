import time

from docx import Document
from docx.opc.oxml import qn
from docx.oxml import OxmlElement
import codecs
from docx.shared import  Pt
from docx.oxml.ns import  qn
from docx.shared import Inches
import re

from crawler.MerInfoCrawler.ChineseSimilartyCaculation import CSV
from crawler.MerInfoCrawler.translate_google import google_translate


def read_txt(path):
    f = codecs.open(path, 'rb', 'utf-8')
    content = f.read()
    f.close()
    return str(content)
def Num_en(sentences):
    sentences = sentences.strip()
    index = 0
    count = 0
    while index < len(sentences):
        while sentences[index] != " ":
            index += 1
            if index == len(sentences):
                break
        count += 1
        if index == len(sentences):
            break
        while sentences[index] == " ":
            index += 1
    return count
def Doc():
    en = []
    cn = []
    cn_pick = []
    right = []

    f = codecs.open('cn.txt', 'rb', 'utf-8')
    content_cn = f.readlines()
    f.close()
    for sentence_cn in content_cn:
        if not sentence_cn.strip() == '':
            sentence_cn = sentence_cn.strip(' ')
            cn.append(sentence_cn)
    f = codecs.open('en.txt', 'rb', 'utf-8')
    content_en = f.readlines()
    f.close()
    flag = -1
    for sentence in content_en:
        if not sentence.strip() == '':
            sentence = sentence.strip(' ')
            if sentence != '|' and sentence != '×':
                flag = flag+1
                if Num_en(sentence) > 5:
                    trans = google_translate(sentence)
                    csv = CSV(trans, 'cn.txt')
                    #print("最相似的文本为：",csv)
                    if csv != 0:
                        en.append(sentence)
                        cn_pick.append(csv)

    for i in range(len(en)):
        print(en[i])
        print("最相似的文本为：", cn_pick[i])
    document = Document('Huawei_cases.docx')
    table = document.tables[0]
    rows = len(table.rows)
    for i in range(len(en)):
        table.add_row()
        table.cell(rows + i, 0).text = cn_pick[i]
        table.cell(rows + i, 1).text = en[i]
    document.save('Huawei_cases.docx')
    '''
    document = Document()
    document.add_heading('华为--新闻', 0)
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    rows = 1
    cols = 2

    table = document.add_table(rows=1, cols=2, style="Table Grid")

    for i in range(rows):
        tr = table.rows[i]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "450")
        trPr.append(trHeight)
    table.cell(0, 0).text = u'中文'
    table.cell(0, 1).text = u'英文'
    now = 1
    for i in range(len(en)):
        table.add_row()
        table.cell(now, 0).text = cn_pick[i]
        table.cell(now, 1).text = en[i]
        now = now + 1
    document.save('Huawei_news.docx')
'''


